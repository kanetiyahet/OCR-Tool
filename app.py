from flask import Flask, render_template, request, jsonify, send_file
import easyocr
import os
import json
import io
import uuid
import base64
import requests
from datetime import datetime
from dotenv import load_dotenv
import google.generativeai as genai
from openai import OpenAI
from sarvamai import SarvamAI

load_dotenv()

SARVAM_API_KEY = os.getenv("SARVAM_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)
openai_client = OpenAI(api_key=OPENAI_API_KEY)
client = SarvamAI(api_subscription_key=SARVAM_API_KEY)

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp', 'tiff'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ── Languages ──────────────────────────────────────────────────────────────
# Full list of EasyOCR-supported language codes
SUPPORTED_LANGUAGES = {
    'en', 'es', 'fr', 'de', 'it', 'pt', 'ru', 'ja', 'ko', 'ch_sim', 'ch_tra',
    'ar', 'hi', 'bn', 'ta', 'te', 'mr', 'ur', 'pa', 'gu', 'kn', 'ml', 'or',
    'ne', 'si', 'th', 'vi', 'id', 'ms', 'tl', 'sw', 'pl', 'nl', 'sv', 'da',
    'no', 'fi', 'hu', 'cs', 'el', 'he', 'fa', 'tr', 'uk'
}

# Languages used for "Auto" detection (most common worldwide + Indian languages)
AUTO_LANGUAGES = [
    'en', 'es', 'fr', 'de', 'it', 'pt', 'ru', 'ja', 'ko', 'ch_sim',
    'ar', 'hi', 'bn', 'ta', 'te', 'mr', 'ur', 'pa', 'gu', 'kn', 'ml'
]

# ── Cache for EasyOCR readers ──────────────────────────────────────────────
reader_cache = {}

def validate_languages(lang_list):
    """Return a list of valid language codes from the given list."""
    if not lang_list:
        return ['en']
    valid = [lang for lang in lang_list if lang in SUPPORTED_LANGUAGES]
    return valid if valid else ['en']

def get_easyocr_reader(lang_codes):
    """Return an EasyOCR reader for the given language code(s)."""
    # Validate and deduplicate
    valid_codes = validate_languages(lang_codes)
    key = tuple(sorted(set(valid_codes)))
    if key not in reader_cache:
        reader_cache[key] = easyocr.Reader(list(key))
    return reader_cache[key]

# ── OCR functions ──────────────────────────────────────────────────────────────

def easyocr_ocr(filepath, language='auto'):
    """If language is 'auto', use the predefined set of languages."""
    if language == 'auto':
        langs = AUTO_LANGUAGES
    else:
        langs = [language] if language else ['en']
    reader = get_easyocr_reader(langs)
    result = reader.readtext(filepath)
    return "\n".join([r[1] for r in result])

def gemini_ocr(filepath):
    try:
        model = genai.GenerativeModel("gemini-2.5-pro")  # use flash to avoid quota issues
        ext = filepath.rsplit('.', 1)[-1].lower()
        mime_map = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
                    'gif': 'image/gif', 'bmp': 'image/bmp', 'webp': 'image/webp'}
        mime_type = mime_map.get(ext, 'image/jpeg')
        with open(filepath, "rb") as f:
            image_bytes = f.read()
        response = model.generate_content([
            "Extract all text from this image. Return only the extracted text, nothing else.",
            {"mime_type": mime_type, "data": image_bytes}
        ])
        return response.text.strip()
    except Exception as e:
        # Fallback to EasyOCR if Gemini fails (e.g., quota exceeded)
        return easyocr_ocr(filepath, 'auto')

def openai_ocr(filepath):
    with open(filepath, "rb") as f:
        image_data = base64.b64encode(f.read()).decode()
    ext = filepath.rsplit('.', 1)[-1].lower()
    mime_map = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
                'gif': 'image/gif', 'bmp': 'image/bmp', 'webp': 'image/webp'}
    mime_type = mime_map.get(ext, 'image/jpeg')
    response = openai_client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": "Extract all text from this image. Return only the extracted text, nothing else."},
                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}}
            ]
        }]
    )
    return response.choices[0].message.content.strip()

def sarvam_ocr(filepath, language):
    ocr_job = client.document_intelligence.create_job(
        language="en-IN",
        output_format="md"
    )
    ocr_job.upload_file(filepath)
    ocr_job.start()
    ocr_job.wait_until_complete()

    zip_file = filepath + ".zip"
    extract_folder = filepath + "_output"
    ocr_job.download_output(zip_file)

    import zipfile
    with zipfile.ZipFile(zip_file, "r") as zip_ref:
        zip_ref.extractall(extract_folder)

    md_file = os.path.join(extract_folder, "document.md")
    if os.path.exists(md_file):
        with open(md_file, "r", encoding="utf-8") as f:
            return f.read()
    return "No OCR output found"

# ── Routes ─────────────────────────────────────────────────────────────────────

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/languages", methods=["GET"])
def languages():
    # Return only supported languages for the dropdown
    return jsonify({code: name for code, name in EASYOCR_LANGUAGE_NAMES.items() if code in SUPPORTED_LANGUAGES})

# Human-readable language names (only for display)
EASYOCR_LANGUAGE_NAMES = {
    'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German',
    'it': 'Italian', 'pt': 'Portuguese', 'ru': 'Russian', 'ja': 'Japanese',
    'ko': 'Korean', 'ch_sim': 'Chinese (Simplified)', 'ch_tra': 'Chinese (Traditional)',
    'ar': 'Arabic', 'hi': 'Hindi', 'bn': 'Bengali', 'ta': 'Tamil',
    'te': 'Telugu', 'mr': 'Marathi', 'ur': 'Urdu', 'pa': 'Punjabi',
    'gu': 'Gujarati', 'kn': 'Kannada', 'ml': 'Malayalam', 'or': 'Odia',
    'ne': 'Nepali', 'si': 'Sinhala', 'th': 'Thai', 'vi': 'Vietnamese',
    'id': 'Indonesian', 'ms': 'Malay', 'tl': 'Tagalog', 'sw': 'Swahili',
    'pl': 'Polish', 'nl': 'Dutch', 'sv': 'Swedish', 'da': 'Danish',
    'no': 'Norwegian', 'fi': 'Finnish', 'hu': 'Hungarian', 'cs': 'Czech',
    'el': 'Greek', 'he': 'Hebrew', 'fa': 'Persian', 'tr': 'Turkish',
    'uk': 'Ukrainian'
}

@app.route("/ocr", methods=["POST"])
def ocr():
    files = request.files.getlist("images")
    if not files or all(f.filename == '' for f in files):
        return jsonify({"error": "No files uploaded"}), 400

    provider = request.form.get("provider", "easyocr")
    language = request.form.get("language", "auto")
    print(f"Provider: {provider}, Language: {language}")

    results = []
    for file in files:
        if file and allowed_file(file.filename):
            ext = file.filename.rsplit('.', 1)[1].lower()
            unique_name = f"{uuid.uuid4().hex}.{ext}"
            filepath = os.path.join(UPLOAD_FOLDER, unique_name)
            file.save(filepath)

            try:
                if provider == "easyocr":
                    text = easyocr_ocr(filepath, language)
                elif provider == "gemini":
                    text = gemini_ocr(filepath)
                elif provider == "openai":
                    text = openai_ocr(filepath)
                elif provider == "sarvam":
                    text = sarvam_ocr(filepath, language)
                else:
                    text = easyocr_ocr(filepath, language)
            except Exception as e:
                text = f"[ERROR] {str(e)}"
            finally:
                if os.path.exists(filepath):
                    os.remove(filepath)

            word_count = len(text.split()) if text.strip() else 0
            char_count = len(text)

            results.append({
                "filename": file.filename,
                "text": text,
                "word_count": word_count,
                "char_count": char_count,
            })
        else:
            results.append({
                "filename": file.filename,
                "error": "Unsupported file type",
                "text": "",
                "word_count": 0,
                "char_count": 0,
            })

    return jsonify({"results": results})

# ── Download routes (unchanged) ──────────────────────────────────────────────
@app.route("/download/txt", methods=["POST"])
def download_txt():
    data = request.get_json()
    text = data.get("text", "")
    buf = io.BytesIO(text.encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="text/plain",
                     as_attachment=True, download_name="ocr_result.txt")

@app.route("/download/json", methods=["POST"])
def download_json():
    data = request.get_json()
    results = data.get("results", [])
    payload = {
        "exported_at": datetime.utcnow().isoformat() + "Z",
        "results": results
    }
    buf = io.BytesIO(json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="application/json",
                     as_attachment=True, download_name="ocr_result.json")

@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
    except ImportError:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500

    data = request.get_json()
    results = data.get("results", [])
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                  textColor=colors.HexColor('#1a73e8'))
    meta_style = ParagraphStyle('Meta', parent=styles['Normal'],
                                 textColor=colors.grey, fontSize=9)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                 leading=16, fontSize=11)
    story = []
    story.append(Paragraph("OCR Results", title_style))
    story.append(Spacer(1, 6*mm))
    for i, r in enumerate(results, 1):
        story.append(Paragraph(f"Image {i}: {r.get('filename','')}", styles['Heading2']))
        story.append(Paragraph(
            f"Words: {r.get('word_count',0)}  |  Characters: {r.get('char_count',0)}",
            meta_style))
        story.append(Spacer(1, 3*mm))
        text = r.get("text", "").replace("\n", "<br/>")
        story.append(Paragraph(text or "<i>(no text detected)</i>", body_style))
        story.append(Spacer(1, 8*mm))
    doc.build(story)
    buf.seek(0)
    return send_file(buf, mimetype="application/pdf",
                     as_attachment=True, download_name="ocr_result.pdf")

@app.route("/download/docx", methods=["POST"])
def download_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        return jsonify({"error": "python-docx not installed. Run: pip install python-docx"}), 500

    data = request.get_json()
    results = data.get("results", [])
    document = Document()
    title = document.add_heading("OCR Results", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
    for i, r in enumerate(results, 1):
        document.add_heading(f"Image {i}: {r.get('filename','')}", level=2)
        meta = document.add_paragraph(
            f"Words: {r.get('word_count',0)}   |   Characters: {r.get('char_count',0)}"
        )
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        text = r.get("text", "") or "(no text detected)"
        p = document.add_paragraph(text)
        p.runs[0].font.size = Pt(11)
        document.add_paragraph()
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="ocr_result.docx")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)